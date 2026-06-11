# -*- coding: utf-8 -*-
"""Generate formatted DOCX report for IS207 project."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"g:\Documents\Web_IS207\Web_IS207\Bao cao IS207.docx"

# Colors — IS207 brown theme
PRIMARY = RGBColor(0x5C, 0x40, 0x33)
ACCENT = RGBColor(0xB8, 0x95, 0x6C)
DARK = RGBColor(0x3D, 0x2E, 0x24)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "5C4033"
TABLE_ALT_BG = "F5F0EB"


def set_cell_shading(cell, fill_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, name="Times New Roman", size=13, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_para(doc, text, style=None, align=None, size=13, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 14, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=sizes.get(level, 13), bold=True, color=PRIMARY)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=13)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=10, color=DARK)
    # light background via shading on paragraph - skip for simplicity


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=11, bold=True, color=WHITE)
        set_cell_shading(hdr_cells[i], TABLE_HEADER_BG)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=11, color=DARK)
            if ri % 2 == 1:
                set_cell_shading(cells[ci], TABLE_ALT_BG)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_note_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Lưu ý: " + text)
    set_run_font(run, size=12, italic=True, color=GRAY)


def build_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    add_para(doc, "TRƯỜNG ĐẠI HỌC ...........................", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True)
    add_para(doc, "KHOA ..........................................", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, "BÁO CÁO MÔN HỌC / ĐỒ ÁN", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, color=PRIMARY, space_after=8)
    add_para(doc, "MÔN: PHÁT TRIỂN ỨNG DỤNG WEB (IS207)", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, color=ACCENT, space_after=20)
    add_para(doc, "ĐỀ TÀI:", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, bold=True, space_after=6)
    add_para(doc, "HỆ THỐNG QUẢN LÝ BÁN HÀNG THỜI TRANG", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, color=PRIMARY, space_after=4)
    add_para(doc, "IS207 FASHION", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True, color=PRIMARY, space_after=4)
    add_para(doc, "(Website khách hàng + Website quản trị)", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, italic=True, color=GRAY, space_after=30)
    fields = [
        ("Sinh viên thực hiện:", "................................"),
        ("MSSV:", "................................"),
        ("Lớp:", "................................"),
        ("Giảng viên hướng dẫn:", "................................"),
    ]
    for label, dots in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + " ")
        set_run_font(r1, size=13, bold=True)
        r2 = p.add_run(dots)
        set_run_font(r2, size=13)
    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, "TP. Hồ Chí Minh, tháng ...... năm 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, space_after=0)
    doc.add_page_break()


def build_toc(doc):
    add_heading(doc, "MỤC LỤC", level=1)
    toc_items = [
        ("", "LỜI CẢM ƠN", ""),
        ("", "DANH MỤC HÌNH VẼ / BẢNG BIỂU", ""),
        ("1", "Giới thiệu chung", "3"),
        ("1.1", "Khái quát về ứng dụng", ""),
        ("1.2", "Mục tiêu và phạm vi", ""),
        ("1.3", "Đối tượng sử dụng", ""),
        ("2", "Công nghệ áp dụng", ""),
        ("2.1", "Kiến trúc hệ thống", ""),
        ("2.2", "Backend", ""),
        ("2.3", "Frontend", ""),
        ("2.4", "Cơ sở dữ liệu và bảo mật", ""),
        ("3", "Giao diện tùy biến (Custom UI & Layouts)", ""),
        ("3.1", "Website khách hàng", ""),
        ("3.2", "Website quản trị", ""),
        ("4", "Các chức năng chính", ""),
        ("4.1", "Chức năng khách hàng", ""),
        ("4.2", "Chức năng quản trị", ""),
        ("4.3", "API backend (REST)", ""),
        ("5", "Khả năng ứng dụng thực tế", ""),
        ("6", "Kiểm thử và triển khai", ""),
        ("6.1", "Kiểm thử", ""),
        ("6.2", "Triển khai", ""),
        ("7", "Kiến thức rút ra", ""),
        ("8", "Hướng phát triển tương lai", ""),
        ("", "TÀI LIỆU THAM KHẢO", ""),
        ("", "PHỤ LỤC", ""),
    ]
    add_table(doc, ["STT", "Nội dung", "Trang"], toc_items, [1.5, 12, 2])
    add_para(doc, "Gợi ý: Trong Word, chọn References → Table of Contents → Automatic Table 1 để cập nhật số trang tự động.", size=11, italic=True, color=GRAY)
    doc.add_page_break()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    build_cover(doc)
    build_toc(doc)

    # === 1. GIỚI THIỆU ===
    add_heading(doc, "1. GIỚI THIỆU CHUNG", level=1)
    add_heading(doc, "1.1. Khái quát về ứng dụng", level=2)
    add_para(doc, "IS207 Fashion là hệ thống bán hàng thời trang trực tuyến gồm ba thành phần trong repository Web_IS207:")
    add_table(doc, ["Thành phần", "Thư mục", "Vai trò"], [
        ["Backend API", "Sale_App/", "Xử lý nghiệp vụ, lưu trữ, xác thực"],
        ["Website khách hàng", "customer/", "Mua sắm, đặt hàng, chính sách"],
        ["Website quản trị", "admin/", "Quản lý SP, đơn hàng, báo cáo"],
    ], [3.5, 3, 8])
    add_para(doc, "Hệ thống không phải ứng dụng Android (không có file APK). Đây là ứng dụng web: giao diện chạy trên trình duyệt, giao tiếp server qua REST API và JSON.")
    add_para(doc, "Luồng hoạt động tổng quát:", bold=True)
    add_code_block(doc, "Trình duyệt (customer/admin) → HTTPS + JSON (JWT)\n        ↓\nSpring Boot API (port 8080 / Railway)\n        ↓\nMySQL (database web_is207)")

    add_heading(doc, "1.2. Mục tiêu và phạm vi", level=2)
    add_para(doc, "Mục tiêu:", bold=True)
    for t in [
        "Xây dựng website bán áo, quần, váy, áo khoác với trải nghiệm mua hàng đầy đủ.",
        "Cung cấp kênh quản trị: sản phẩm, đơn hàng, thống kê từ dữ liệu thật.",
        "Áp dụng kiến trúc tách lớp: frontend tĩnh + backend API + database.",
    ]:
        add_bullet(doc, t)
    add_para(doc, "Phạm vi đã triển khai:", bold=True)
    for t in [
        "Đăng ký/đăng nhập JWT; quản lý SP (CRUD, import CSV/Excel, upload ảnh).",
        "Giỏ hàng, đặt hàng, 5 phương thức thanh toán (mô phỏng).",
        "Quản lý trạng thái đơn; báo cáo SVG trên admin.",
        "Chính sách, hướng dẫn, chat auto-reply.",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "1.3. Đối tượng sử dụng", level=2)
    add_table(doc, ["Đối tượng", "Vai trò"], [
        ["Khách hàng", "Xem SP, lọc, yêu thích, đặt hàng, theo dõi đơn"],
        ["Quản trị viên", "Quản lý danh mục, đơn hàng, báo cáo, cập nhật trạng thái"],
    ], [4, 11])

    # === 2. CÔNG NGHỆ ===
    add_heading(doc, "2. CÔNG NGHỆ ÁP DỤNG", level=1)
    add_heading(doc, "2.1. Kiến trúc hệ thống", level=2)
    add_bullet(doc, "Kiến trúc 3 tầng: Presentation (HTML/CSS/JS) — Application (Spring Boot) — Data (MySQL/JPA)")
    add_bullet(doc, "Giao tiếp: REST + JSON qua HTTP/HTTPS")
    add_bullet(doc, "Xác thực: JWT, token lưu localStorage")

    add_heading(doc, "2.2. Backend", level=2)
    add_table(doc, ["Công nghệ", "Phiên bản/Ghi chú", "Mục đích"], [
        ["Java", "17", "Ngôn ngữ backend"],
        ["Spring Boot", "4.0.6", "Framework API"],
        ["Spring Data JPA", "—", "ORM"],
        ["Spring Security", "—", "JWT, phân quyền"],
        ["Spring Mail", "—", "Quên mật khẩu"],
        ["MySQL Connector", "—", "Kết nối DB"],
        ["JJWT", "0.12.6", "JWT"],
        ["Docker", "Dockerfile", "Deploy cloud"],
    ], [3.5, 4, 6])

    add_heading(doc, "2.3. Frontend", level=2)
    add_table(doc, ["Công nghệ", "Ghi chú", "Mục đích"], [
        ["HTML5 / CSS3", "design-tokens.css", "Giao diện responsive"],
        ["JavaScript ES6+", "Không React/Vue", "Logic UI, API"],
        ["Font Awesome 6.4", "CDN", "Icon"],
        ["SheetJS (xlsx)", "admin", "Import Excel"],
        ["Fetch API", "api.js", "HTTP client"],
    ], [3.5, 5, 5])

    add_heading(doc, "2.4. Cơ sở dữ liệu và bảo mật", level=2)
    for t in [
        "MySQL — database web_is207",
        "JWT — Authorization: Bearer <token>",
        "Mật khẩu hash (BCrypt)",
        "Upload ảnh — thư mục uploads/",
        "CORS — SecurityConfig.java",
    ]:
        add_bullet(doc, t)

    # === 3. UI ===
    add_heading(doc, "3. GIAO DIỆN TÙY BIẾN (CUSTOM UI & LAYOUTS)", level=1)
    add_note_box(doc, "Đề cương mobile ghi \"Multiple Activities\". Dự án web tương đương nhiều màn hình/views — không có Activity Android.")
    add_heading(doc, "3.1. Website khách hàng", level=2)
    for t in [
        "Design tokens màu nâu/be; header cố định.",
        "Trang chủ: carousel, flash sale, SP nổi bật.",
        "Lọc size/màu/giá/giới tính theo danh mục.",
        "Modal: chi tiết SP, giỏ, checkout, chính sách.",
        "Thanh toán thẻ + OTP 3-D Secure (mô phỏng); chat widget.",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "3.2. Website quản trị", level=2)
    add_table(doc, ["Trang", "Nội dung UI"], [
        ["dashboard", "KPI, biểu đồ 7 ngày, đơn gần đây, cảnh báo tồn"],
        ["products", "Bảng SP, CRUD, import CSV/Excel"],
        ["orders", "Bảng đơn, modal chi tiết, đổi trạng thái"],
        ["customers", "Tổng hợp từ đơn, lịch sử mua"],
        ["reports", "KPI kỳ, 6 biểu đồ SVG"],
    ], [3, 11])

    # === 4. CHỨC NĂNG ===
    add_heading(doc, "4. CÁC CHỨC NĂNG CHÍNH", level=1)
    add_heading(doc, "4.1. Chức năng phía khách hàng", level=2)
    cust_funcs = [
        ("1", "Đăng ký/đăng nhập", "JWT customer"),
        ("2", "Check trùng email/SĐT", "API check-availability"),
        ("3", "Quên mật khẩu", "Mã 6 số"),
        ("4", "Xem & tìm SP", "/api/products"),
        ("5", "Lọc & sắp xếp", "Size, màu, giá, giới tính"),
        ("6", "Khuyến mãi", "Banner, flash sale"),
        ("7", "Chi tiết SP", "Tab mô tả/thông số/bảo hành"),
        ("8", "Wishlist", "localStorage"),
        ("9", "Giỏ hàng", "localStorage"),
        ("10", "Checkout", "Form + mã giảm giá"),
        ("11", "Thanh toán", "COD, VietQR, Momo, ZaloPay, thẻ+OTP"),
        ("12", "Đơn của tôi", "Khi đăng nhập"),
        ("13", "Chính sách & chat", "Modal + auto-reply"),
    ]
    add_table(doc, ["STT", "Chức năng", "Mô tả"], cust_funcs, [1, 4.5, 8])

    add_heading(doc, "4.2. Chức năng phía quản trị", level=2)
    admin_funcs = [
        ("1", "Đăng nhập admin", "JWT admin"),
        ("2", "Dashboard", "KPI, so sánh hôm qua"),
        ("3", "Quản lý SP", "CRUD, import, upload"),
        ("4", "Quản lý đơn", "Filter, modal, đổi trạng thái"),
        ("5", "Khách hàng", "Tổng hợp từ đơn"),
        ("6", "Báo cáo", "Biểu đồ SVG, KPI kỳ"),
        ("7", "Thông báo", "Đơn pending"),
    ]
    add_table(doc, ["STT", "Chức năng", "Mô tả"], admin_funcs, [1, 4.5, 8])
    add_para(doc, "Lưu ý: Trang Bán hàng (POS) đã loại bỏ — admin không bán trực tiếp.", italic=True, color=GRAY)

    add_heading(doc, "4.3. API backend (REST)", level=2)
    apis = [
        ("Auth", "POST /api/auth/login", "Đăng nhập"),
        ("Auth", "GET /api/auth/check-availability", "Check phone/email"),
        ("Auth", "POST /api/auth/register", "Đăng ký"),
        ("Products", "GET/POST /api/products", "Danh sách/Tạo"),
        ("Orders", "POST /api/orders", "Tạo đơn"),
        ("Orders", "PUT .../status", "Cập nhật trạng thái"),
        ("Upload", "POST /api/uploads", "Upload ảnh"),
    ]
    add_table(doc, ["Nhóm", "Endpoint", "Mô tả"], apis, [2.5, 5.5, 5])
    add_para(doc, "Trạng thái đơn: pending → shipping → completed / cancelled")

    # === 5. THỰC TẾ ===
    add_heading(doc, "5. KHẢ NĂNG ỨNG DỤNG TRONG THỰC TẾ", level=1)
    add_table(doc, ["Bối cảnh", "Ứng dụng"], [
        ["Shop thời trang nhỏ", "Bán online, admin xử lý đơn"],
        ["Showroom + online", "Khách đặt web, cửa hàng chuẩn bị hàng"],
        ["Đồ án / demo", "Luồng e-commerce đầy đủ"],
    ], [5, 9])
    add_para(doc, "Giá trị: giảm nhập đơn thủ công; báo cáo hỗ trợ nhập hàng; phân quyền bảo vệ dữ liệu.")

    # === 6. KIỂM THỬ & TRIỂN KHAI ===
    add_heading(doc, "6. KIỂM THỬ VÀ TRIỂN KHAI", level=1)
    add_note_box(doc, "Dự án không build APK. Phần này mô tả kiểm thử và triển khai web + API.")
    add_heading(doc, "6.1. Kiểm thử", level=2)
    add_para(doc, "Môi trường local:", bold=True)
    add_code_block(doc, "cd Sale_App\n./mvnw.cmd spring-boot:run\n# Mở customer/index.html và admin/index.html qua Live Server")
    add_para(doc, "Kịch bản test:", bold=True)
    tests = [
        ("Đăng ký + check trùng", "Báo lỗi/OK realtime"),
        ("Checkout COD", "Admin thấy đơn mới"),
        ("Thanh toán thẻ + OTP", "Luhn + modal OTP"),
        ("Admin đổi trạng thái", "pending → completed"),
        ("Import CSV", "Thêm hàng loạt SP"),
    ]
    add_table(doc, ["Kịch bản", "Kết quả mong đợi"], tests, [5, 9])
    add_para(doc, "Build JAR:", bold=True)
    add_code_block(doc, "cd Sale_App && ./mvnw.cmd -DskipTests package")

    add_heading(doc, "6.2. Triển khai", level=2)
    add_table(doc, ["Thành phần", "Nền tảng"], [
        ["Backend + MySQL", "Railway"],
        ["Frontend customer", "Vercel / Netlify"],
        ["Frontend admin", "Vercel / Netlify"],
    ], [5, 8])
    add_para(doc, "Sau deploy, set backend URL trong Console trình duyệt:")
    add_code_block(doc, 'localStorage.setItem("backend_base_url", "https://<railway-domain>");\nlocation.reload();')

    # === 7. KIẾN THỨC ===
    add_heading(doc, "7. KIẾN THỨC RÚT RA", level=1)
    knowledge = [
        "Kiến trúc tách frontend–backend: deploy linh hoạt (Vercel + Railway).",
        "Spring Boot & JPA: entity, repository, service, controller.",
        "Bảo mật web: JWT, Spring Security, BCrypt.",
        "JavaScript client: state, filter, validation (Luhn), async API.",
        "UX e-commerce: checkout, thanh toán, chính sách.",
        "Báo cáo: KPI và biểu đồ SVG từ dữ liệu thật.",
        "DevOps: Maven, Docker, biến môi trường, CORS.",
        "Git: branch, commit có ý nghĩa.",
    ]
    for i, k in enumerate(knowledge, 1):
        add_para(doc, f"{i}. {k}")

    # === 8. HƯỚNG PHÁT TRIỂN ===
    add_heading(doc, "8. HƯỚNG PHÁT TRIỂN TƯƠNG LAI", level=1)
    future = [
        ("Cổng thanh toán thật", "VNPay, Momo, ZaloPay API"),
        ("CDN ảnh", "Cloudinary / S3"),
        ("CORS production", "Domain Vercel cố định"),
        ("Real-time", "WebSocket cho đơn mới"),
        ("Mobile app", "React Native / Flutter + API"),
        ("Voucher backend", "Mã giảm giá lưu DB"),
        ("Export báo cáo", "PDF/Excel"),
        ("Unit test", "JUnit, API test"),
    ]
    add_table(doc, ["Hướng", "Mô tả"], future, [4.5, 9])

    # === TÀI LIỆU THAM KHẢO ===
    add_heading(doc, "TÀI LIỆU THAM KHẢO", level=1)
    refs = [
        "Spring Boot Documentation — https://docs.spring.io/spring-boot/",
        "Spring Security — https://docs.spring.io/spring-security/reference/",
        "Spring Data JPA — https://docs.spring.io/spring-data/jpa/",
        "MDN Web Docs — https://developer.mozilla.org/",
        "JWT — https://jwt.io/introduction",
        "MySQL 8.0 — https://dev.mysql.com/doc/",
        "REST API — https://restfulapi.net/",
        "Railway — https://docs.railway.app/",
        "Vercel — https://vercel.com/docs",
        "OWASP — https://owasp.org/",
    ]
    for i, r in enumerate(refs, 1):
        add_para(doc, f"[{i}] {r}", size=12)

    # === PHỤ LỤC ===
    add_heading(doc, "PHỤ LỤC", level=1)
    for t in [
        "Phụ lục A: Sơ đồ use case",
        "Phụ lục B: Sơ đồ luồng đặt hàng",
        "Phụ lục C: Ảnh chụp màn hình",
        "Phụ lục D: Danh sách API",
        "Phụ lục E: Hướng dẫn cài đặt (DEPLOY-BACKEND.md)",
    ]:
        add_bullet(doc, t)

    doc.add_paragraph()
    add_para(doc, "— Báo cáo soạn từ mã nguồn Web_IS207 (nhánh feature/full-project-2026-05) —", size=11, italic=True, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
