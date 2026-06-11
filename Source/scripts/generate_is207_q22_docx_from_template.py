from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@dataclass(frozen=True)
class Figure:
    caption: str
    path: Path
    width_inches: float = 6.5


TEMPLATE_PATH = Path(r"g:\Template_Do_An_Mon_Hoc_VN (1).docx")
OUTPUT_PATH = Path(r"g:\Documents\Web_IS207\Web_IS207\BaoCao_IS207_Q22_v4.docx")

ASSETS_DIR = Path(r"C:\Users\TEC\.cursor\projects\g-Documents-Web-IS207\assets")


def clear_document_keep_styles(doc: Document) -> None:
    body = doc._element.body  # noqa: SLF001
    # Keep section properties (<w:sectPr>) so python-docx can compute page width.
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text)
    # Try to use Word heading styles if present; otherwise fall back to Normal with bold.
    style = f"Heading {level}"
    try:
        p.style = style
    except Exception:
        run = p.runs[0] if p.runs else p.add_run(text)
        run.bold = True
        run.font.size = Pt(14 if level == 1 else 13)


def add_paragraph(doc: Document, text: str, *, align: int | None = None) -> None:
    p = doc.add_paragraph(text)
    if align is not None:
        p.alignment = align


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_toc_field(doc: Document) -> None:
    """
    Insert a Word TOC field. Word renders it after user updates fields (F9).
    """
    p = doc.add_paragraph()
    r = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \\o "1-3" \\h \\z \\u'

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_text = OxmlElement("w:t")
    fld_text.text = "Mục lục sẽ hiển thị sau khi cập nhật (References → Update Table)."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r._r.append(fld_begin)  # noqa: SLF001
    r._r.append(instr)  # noqa: SLF001
    r._r.append(fld_sep)  # noqa: SLF001
    r._r.append(fld_text)  # noqa: SLF001
    r._r.append(fld_end)  # noqa: SLF001


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(it)
        try:
            p.style = "List Bullet"
        except Exception:
            pass


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v


def add_figure(doc: Document, fig: Figure, idx: int) -> None:
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_img.add_run()
    run.add_picture(str(fig.path), width=Inches(fig.width_inches))

    p_cap = doc.add_paragraph(f"Hình {idx}. {fig.caption}")
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p_cap.style = "Caption"
    except Exception:
        pass


def main() -> None:
    if not TEMPLATE_PATH.exists():
        raise SystemExit(f"Template not found: {TEMPLATE_PATH}")

    # Selected screenshots (customer + admin + deploy). Keep concise but representative.
    figures: list[Figure] = [
        Figure(
            caption="Giao diện trang chủ website khách hàng (IS207 Fashion).",
            path=ASSETS_DIR
            / "c__Users_TEC_AppData_Roaming_Cursor_User_workspaceStorage_72067249eedbbc3ce769303f1f924858_images_Screenshot_2026-05-28_174124-9610521d-54dd-4c2a-8fba-370c70b67b1d.png",
        ),
        Figure(
            caption="Trang danh sách sản phẩm với bộ lọc và giỏ hàng (customer).",
            path=ASSETS_DIR
            / "c__Users_TEC_AppData_Roaming_Cursor_User_workspaceStorage_72067249eedbbc3ce769303f1f924858_images_Screenshot_2026-05-28_174319-f7bebb5a-0831-4c74-a581-619aab03822a.png",
        ),
        Figure(
            caption="Trang thông tin liên hệ (modal) của website khách hàng.",
            path=ASSETS_DIR
            / "c__Users_TEC_AppData_Roaming_Cursor_User_workspaceStorage_72067249eedbbc3ce769303f1f924858_images_Screenshot_2026-05-28_174244-eb4a8b84-a5a4-4bdd-a179-aa6e798efbb9.png",
        ),
        Figure(
            caption="Dashboard quản trị (admin) tổng quan KPI và biểu đồ.",
            path=ASSETS_DIR
            / "c__Users_TEC_AppData_Roaming_Cursor_User_workspaceStorage_72067249eedbbc3ce769303f1f924858_images_Screenshot_2026-05-28_174750-26204d8d-19e2-4f54-aef3-7cc1d4062f8e.png",
        ),
        Figure(
            caption="Trang quản lý sản phẩm (admin) với danh sách sản phẩm và thao tác CRUD.",
            path=ASSETS_DIR
            / "c__Users_TEC_AppData_Roaming_Cursor_User_workspaceStorage_72067249eedbbc3ce769303f1f924858_images_Screenshot_2026-05-28_174755-4d6f8bd2-59d4-4821-bdc2-820de775d155.png",
        ),
        Figure(
            caption="Trang báo cáo & thống kê (admin) với biểu đồ doanh thu và KPI theo kỳ.",
            path=ASSETS_DIR
            / "c__Users_TEC_AppData_Roaming_Cursor_User_workspaceStorage_72067249eedbbc3ce769303f1f924858_images_Screenshot_2026-05-28_174820-7874c3ce-bf72-423a-8727-a38f662b8a1f.png",
        ),
        Figure(
            caption="Triển khai backend + MySQL trên Railway (minh hoạ quá trình deploy).",
            path=ASSETS_DIR
            / "c__Users_TEC_AppData_Roaming_Cursor_User_workspaceStorage_72067249eedbbc3ce769303f1f924858_images_image-5002585e-27d4-45de-bf06-118bb417e704.png",
        ),
    ]

    # Validate figure paths exist (skip missing silently so script is robust).
    figures = [f for f in figures if f.path.exists()]

    doc = Document(str(TEMPLATE_PATH))
    clear_document_keep_styles(doc)

    # No cover page per request (template cover removed).
    add_heading(doc, "MỤC LỤC", 1)
    add_paragraph(doc, "Trong Word: References → Table of Contents → chọn mẫu, hoặc nhấn F9 để cập nhật.")
    add_toc_field(doc)
    add_page_break(doc)

    add_heading(doc, "DANH SÁCH HÌNH, BẢNG", 1)
    for i, fig in enumerate(figures, start=1):
        add_paragraph(doc, f"Hình {i}. {fig.caption}")
    add_paragraph(doc, "Bảng 1. Mô tả các bảng dữ liệu chính (users, products, orders, order_items, password_reset_tokens).")
    add_paragraph(doc, "Bảng 2. Danh sách API endpoints chính.")
    add_paragraph(doc, "Bảng 3. Các chức năng chính và mô tả triển khai.")
    add_page_break(doc)

    add_heading(doc, "TÓM TẮT", 1)
    add_paragraph(
        doc,
        "Đồ án IS207 Fashion là hệ thống thương mại điện tử thời trang gồm website khách hàng và trang quản trị, "
        "kết nối với Backend REST API (Spring Boot) và CSDL MySQL. Hệ thống hỗ trợ các nghiệp vụ chính: đăng ký/đăng nhập, "
        "duyệt sản phẩm – lọc theo thuộc tính, giỏ hàng, đặt hàng, quản lý đơn hàng và quản lý sản phẩm trên admin. "
        "Ứng dụng triển khai theo mô hình client–server, xác thực bằng JWT, dữ liệu được lưu trữ tập trung trên MySQL, "
        "hình ảnh sản phẩm lưu trên Cloudinary và chức năng quên mật khẩu gửi email qua Resend (HTTPS API). "
        "Kết quả đạt được là một hệ thống demo chạy online, giao diện thân thiện, luồng nghiệp vụ đầy đủ, và có tài liệu hướng dẫn triển khai.",
    )

    add_heading(doc, "Chương I. TỔNG QUAN.", 1)
    add_heading(doc, "1. Giới thiệu đề tài", 2)
    add_paragraph(
        doc,
        "Trong bối cảnh mua sắm online ngày càng phổ biến, một website bán hàng thời trang cần đáp ứng các yếu tố: "
        "giao diện dễ dùng, tìm kiếm/lọc nhanh, đặt hàng thuận tiện, quản trị sản phẩm/đơn hàng hiệu quả và đảm bảo an toàn đăng nhập. "
        "Nhóm thực hiện đề tài IS207 Fashion nhằm xây dựng hệ thống web e-commerce theo mô hình thực tế, áp dụng kiến thức môn học IS207 "
        "về kiến trúc hệ thống, triển khai dịch vụ, bảo mật và tích hợp.",
    )
    add_heading(doc, "2. Mục tiêu và phạm vi", 2)
    add_heading(doc, "2.1. Mục tiêu", 3)
    add_bullets(
        doc,
        [
            "Xây dựng hệ thống web bán hàng thời trang đầy đủ luồng từ khách hàng đến quản trị.",
            "Thiết kế REST API chuẩn hóa dữ liệu, kết nối MySQL, phân quyền bằng JWT.",
            "Triển khai chạy online (Backend + DB + Frontend) và có hướng dẫn deploy.",
        ],
    )
    add_heading(doc, "2.2. Phạm vi", 3)
    add_bullets(
        doc,
        [
            "Customer site: xem sản phẩm, lọc, chi tiết sản phẩm, giỏ hàng, checkout, đặt hàng, quên mật khẩu.",
            "Admin site: quản lý sản phẩm, quản lý đơn hàng, dashboard thống kê, báo cáo.",
            "Không triển khai thanh toán thật (thanh toán thẻ/OTP là mô phỏng phục vụ demo).",
        ],
    )
    add_heading(doc, "3. Công nghệ sử dụng", 2)
    add_heading(doc, "3.1. Backend (Spring Boot)", 3)
    add_table(
        doc,
        headers=["Công nghệ", "Vai trò"],
        rows=[
            ["Java 17", "Ngôn ngữ backend"],
            ["Spring Boot 4.0.6", "Framework xây dựng REST API, cấu hình theo profile/env"],
            ["Spring Web (starter-web)", "Controller, REST endpoints, JSON"],
            ["Spring Security (starter-security)", "Phân quyền endpoint, JWT filter, bảo vệ API"],
            ["Spring Validation (starter-validation)", "Validate request body/params"],
            ["Spring Data JPA (starter-data-jpa)", "ORM thao tác MySQL qua entity/repository"],
            ["Hibernate ORM", "Cơ chế mapping object-relational (JPA provider)"],
            ["MySQL Connector/J", "Driver JDBC kết nối MySQL"],
            ["JJWT 0.12.6", "Tạo và kiểm tra JWT (access token)"],
            ["Jackson Databind", "Serialize/deserialize JSON"],
            ["Spring Mail (starter-mail)", "Tầng email (fallback SMTP; hiện dùng API email)"],
            ["Docker (Dockerfile)", "Đóng gói ứng dụng để deploy trên Railway"],
            ["railway.toml", "Khai báo build & healthcheck cho Railway"],
        ],
    )

    add_heading(doc, "3.2. Database (MySQL)", 3)
    add_bullets(
        doc,
        [
            "MySQL (deploy trên Railway).",
            "Thiết kế dữ liệu quan hệ gồm: users, products, orders, order_items, password_reset_tokens.",
            "ORM qua JPA/Hibernate, ddl-auto=update phù hợp môi trường demo.",
        ],
    )

    add_heading(doc, "3.3. Frontend (Customer/Admin)", 3)
    add_table(
        doc,
        headers=["Công nghệ", "Vai trò"],
        rows=[
            ["HTML5", "Cấu trúc trang, semantic sections, modal"],
            ["CSS3", "Giao diện, responsive, design tokens"],
            ["JavaScript (ES6+)", "Xử lý UI, gọi API, state localStorage"],
            ["Fetch API", "Gọi REST API (JSON)"],
            ["localStorage", "Lưu token đăng nhập, giỏ hàng, wishlist, backend_base_url"],
            ["SVG (admin reports)", "Vẽ biểu đồ thống kê (line/bar/donut) không cần thư viện nặng"],
            ["SheetJS (xlsx) (admin)", "Import Excel/CSV sản phẩm (phía admin)"],
        ],
    )

    add_heading(doc, "3.4. Hạ tầng, dịch vụ & tích hợp", 3)
    add_table(
        doc,
        headers=["Thành phần", "Dịch vụ", "Mục đích"],
        rows=[
            ["Backend API", "Railway", "Deploy Spring Boot + cấu hình env + domain public"],
            ["Database", "Railway MySQL", "Lưu dữ liệu tập trung"],
            ["Frontend customer", "Vercel", "Host static site, gọi API Railway"],
            ["Ảnh sản phẩm", "Cloudinary", "Lưu ảnh qua URL public (CDN)"],
            ["Email reset password", "Resend (HTTPS API)", "Gửi OTP, tránh SMTP bị chặn trên Railway"],
        ],
    )

    add_heading(doc, "3.5. Công cụ phát triển", 3)
    add_bullets(
        doc,
        [
            "Git/GitHub: quản lý phiên bản và đồng bộ mã nguồn.",
            "Browser DevTools: debug CORS, token, request/response, network.",
            "Postman (tuỳ chọn): test API endpoints nhanh.",
        ],
    )
    add_heading(doc, "4. Cơ sở lý thuyết", 2)
    add_heading(doc, "4.1. RESTful API", 3)
    add_paragraph(
        doc,
        "REST (Representational State Transfer) là phong cách thiết kế dịch vụ web dựa trên tài nguyên (resource). "
        "Ứng dụng giao tiếp qua HTTP/HTTPS, dữ liệu trao đổi chủ yếu là JSON, sử dụng các phương thức GET/POST/PUT/DELETE "
        "tương ứng với thao tác đọc/tạo/cập nhật/xóa.",
    )
    add_heading(doc, "4.2. JWT (JSON Web Token)", 3)
    add_paragraph(
        doc,
        "JWT là chuẩn token dạng chuỗi, gồm 3 phần (header.payload.signature). Backend ký token bằng JWT_SECRET và "
        "frontend gửi token trong header Authorization: Bearer <token>. JWT giúp hệ thống xác thực stateless, "
        "phù hợp mô hình client–server và dễ triển khai trên cloud.",
    )
    add_heading(doc, "4.3. ORM với JPA/Hibernate", 3)
    add_paragraph(
        doc,
        "JPA/Hibernate ánh xạ (mapping) object Java sang bảng dữ liệu quan hệ trong MySQL, giúp giảm viết SQL thủ công "
        "và quản lý entity, repository thuận tiện. Project sử dụng Spring Data JPA để thao tác CRUD và truy vấn.",
    )
    add_heading(doc, "4.4. CORS", 3)
    add_paragraph(
        doc,
        "CORS (Cross-Origin Resource Sharing) là cơ chế trình duyệt kiểm soát việc trang web từ một origin gọi tài nguyên "
        "ở origin khác. Khi deploy frontend trên Vercel và backend trên Railway, cần cấu hình CORS trên backend để cho phép "
        "domain Vercel truy cập API.",
    )
    add_heading(doc, "4.5. Cloudinary", 3)
    add_paragraph(
        doc,
        "Cloudinary là dịch vụ lưu trữ và phân phối ảnh qua URL public (CDN). Thay vì lưu ảnh trên ổ đĩa server (dễ mất khi redeploy), "
        "project lưu URL ảnh Cloudinary trong bảng products để đảm bảo ảnh ổn định trên môi trường cloud.",
    )
    add_heading(doc, "4.6. Resend (Email API)", 3)
    add_paragraph(
        doc,
        "Resend cung cấp API gửi email qua HTTPS. Trong môi trường PaaS như Railway, SMTP có thể bị chặn hoặc không ổn định, "
        "nên sử dụng Resend giúp gửi mail quên mật khẩu ổn định hơn.",
    )
    add_heading(doc, "4.7. Docker & triển khai PaaS", 3)
    add_paragraph(
        doc,
        "Docker đóng gói ứng dụng cùng runtime để chạy nhất quán giữa local và cloud. Backend Spring Boot được build thành JAR "
        "và chạy trong container trên Railway; các thông số cấu hình (DB, JWT, CORS) được đưa vào qua biến môi trường.",
    )
    add_heading(doc, "4.8. Bảo mật cơ bản (Authentication/Authorization)", 3)
    add_paragraph(
        doc,
        "Hệ thống áp dụng xác thực JWT và phân quyền theo vai trò (customer/admin). Endpoint nhạy cảm như "
        "tạo/sửa/xóa sản phẩm và cập nhật trạng thái đơn hàng yêu cầu quyền admin. Mật khẩu lưu dạng hash (BCrypt).",
    )

    add_heading(doc, "Chương II. THIẾT KẾ HỆ THỐNG.", 1)
    add_heading(doc, "1. Kiến trúc tổng thể", 2)
    add_paragraph(
        doc,
        "Hệ thống theo mô hình 3 lớp: Presentation (Frontend) – Application (Backend API) – Data (Database).",
    )
    add_bullets(
        doc,
        [
            "Presentation (Frontend): Customer + Admin (static web trên Vercel hoặc chạy local).",
            "Application (Backend API): Spring Boot (Railway).",
            "Data (Database): MySQL (Railway).",
        ],
    )
    add_paragraph(
        doc,
        "Luồng tổng quát: Người dùng thao tác trên Customer/Admin → Frontend gọi REST API → Backend xử lý nghiệp vụ → "
        "lưu/đọc MySQL → trả JSON cho Frontend hiển thị. Ảnh sản phẩm lưu Cloudinary, Backend lưu URL ảnh vào bảng products. "
        "Quên mật khẩu: Backend tạo mã OTP và gửi email qua Resend.",
    )
    # Insert selected figures relevant to design/deploy.
    fig_idx = 1
    for fig in figures:
        add_figure(doc, fig, fig_idx)
        fig_idx += 1

    add_heading(doc, "2. Thiết kế cơ sở dữ liệu (mô tả)", 2)
    add_table(
        doc,
        headers=["Bảng", "Mô tả"],
        rows=[
            ["users", "Lưu tài khoản customer/admin, thông tin (username/phone/email/fullName/address), mật khẩu hash, role."],
            ["products", "name, category, sku (unique), price, brand, material, colors (json), sizes (json), stockBySize (json), image (URL)."],
            ["orders", "Thông tin người nhận, trạng thái (pending/shipping/completed/cancelled), tổng tiền, ghi chú thanh toán."],
            ["order_items", "Danh sách sản phẩm trong đơn: productId, sku, name, price, qty, size, color."],
            ["password_reset_tokens", "Token 6 số, userId, thời gian tạo/hết hạn, used."],
        ],
    )
    add_heading(doc, "3. Thiết kế API (tóm tắt)", 2)
    add_table(
        doc,
        headers=["Nhóm", "Endpoint tiêu biểu"],
        rows=[
            ["Auth", "POST /api/auth/login; POST /api/auth/register; GET /api/auth/check-availability; POST /api/auth/forgot-password; POST /api/auth/reset-password"],
            ["User", "GET /api/me; PUT /api/me"],
            ["Products", "GET /api/products (public); POST/PUT/DELETE /api/products (admin)"],
            ["Orders", "POST /api/orders (customer); GET /api/orders, GET /api/orders/{id} (authenticated); PUT /api/orders/{id}/status (admin)"],
            ["Uploads", "POST /api/uploads (admin) — chỉ dùng local/demo"],
        ],
    )
    add_heading(doc, "4. Thiết kế giao diện", 2)
    add_heading(doc, "4.1. Customer", 3)
    add_bullets(
        doc,
        [
            "Trang chủ: banner/collection, sản phẩm nổi bật.",
            "Danh sách sản phẩm: lọc theo size/màu/giá/giới tính.",
            "Chi tiết sản phẩm: mô tả, thông số, chính sách.",
            "Cart/Checkout: form thông tin + phương thức thanh toán mô phỏng.",
        ],
    )
    add_heading(doc, "4.2. Admin", 3)
    add_bullets(
        doc,
        [
            "Dashboard KPI + chart.",
            "Products CRUD.",
            "Orders list + modal chi tiết + cập nhật trạng thái.",
            "Reports: thống kê theo mốc thời gian (7/30/90/365/custom).",
        ],
    )

    add_heading(doc, "Chương III. TRIỂN KHAI HỆ THỐNG.", 1)
    add_heading(doc, "1. Triển khai Backend (Railway)", 2)
    add_paragraph(doc, "Backend đóng gói bằng Docker, chạy Spring Boot trên Railway.")
    add_bullets(
        doc,
        [
            "Kết nối MySQL nội bộ Railway mysql.railway.internal:3306.",
            "Biến môi trường quan trọng: SPRING_PROFILES_ACTIVE=cloud; SPRING_DATASOURCE_URL/USERNAME/PASSWORD; JWT_SECRET; APP_CORS_ORIGINS.",
            "Email: RESEND_API_KEY, MAIL_ENABLED, MAIL_FROM, MAIL_FROM_NAME.",
        ],
    )
    add_heading(doc, "2. Triển khai Frontend (Vercel)", 2)
    add_bullets(
        doc,
        [
            "Deploy repo FE (customer) lên Vercel dạng static site.",
            "Frontend trỏ backend Railway bằng cấu hình config.js hoặc localStorage.",
        ],
    )
    add_heading(doc, "3. Chức năng chính (mô tả triển khai)", 2)
    add_table(
        doc,
        headers=["Chức năng", "Mô tả"],
        rows=[
            ["Đăng nhập/đăng ký", "Backend tạo JWT, frontend lưu token (localStorage) và gửi Authorization header khi gọi API."],
            ["Sản phẩm", "Admin tạo/sửa/xoá; customer xem danh sách và chi tiết. Ảnh sản phẩm dùng URL Cloudinary."],
            ["Đặt hàng", "Customer tạo order → backend lưu orders/order_items. Admin xem và cập nhật trạng thái."],
            ["Báo cáo/Thống kê", "Admin tổng hợp đơn theo ngày/trạng thái, hiển thị biểu đồ."],
            ["Quên mật khẩu", "Backend tạo mã 6 số, lưu token + expiry, gửi email qua Resend. Người dùng nhập mã để đặt mật khẩu mới."],
        ],
    )
    add_heading(doc, "4. Kiểm thử & đánh giá", 2)
    add_bullets(
        doc,
        [
            "Kiểm thử API bằng trình duyệt/Postman: /api/products, /api/auth/login, /api/orders.",
            "Kiểm thử CORS bằng FE Vercel gọi API Railway.",
            "Kiểm thử quên mật khẩu: kiểm tra email nhận mã, token hết hạn, token dùng 1 lần.",
        ],
    )

    add_heading(doc, "Chương IV. KẾT LUẬN.", 1)
    add_paragraph(
        doc,
        "Nhóm đã hoàn thiện hệ thống web thương mại điện tử thời trang gồm backend REST API, database MySQL và frontend khách hàng/quản trị. "
        "Các chức năng chính hoạt động ổn định: đăng nhập JWT, CRUD sản phẩm, đặt hàng, quản trị đơn hàng, báo cáo và quên mật khẩu qua email. "
        "Hạn chế hiện tại là chưa tích hợp cổng thanh toán thật và một số tính năng nâng cao (tìm kiếm full-text, phân quyền chi tiết hơn, "
        "tối ưu hiệu năng cho dữ liệu lớn). Hướng phát triển tiếp theo gồm tích hợp thanh toán thật, tối ưu UI/UX, triển khai CI/CD, "
        "và nâng cấp hệ thống log/monitoring.",
    )

    add_heading(doc, "NGUỒN THAM KHẢO", 1)
    add_bullets(
        doc,
        [
            "Spring Boot Documentation (Spring Security, Spring Data JPA, Mail)",
            "JWT.io – JSON Web Tokens",
            "MySQL Documentation",
            "Railway Documentation (deploy + environment variables)",
            "Vercel Documentation (static site deploy)",
            "Cloudinary Documentation (image hosting)",
            "Resend API Documentation (send email)",
        ],
    )

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()

